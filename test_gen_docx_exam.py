import io
from docx import Document
from docx.shared import Pt , Inches
from docx.oxml.ns import qn
from pystrich.code128 import Code128Encoder


TEST_DATA = ("准考号:  ","姓　名:  ","性　别:  ","考　点:  ","报名点:  ")
STUD_DATA = ('18251402',"李文娟",'男','泗县一中','01中学')

def chg_font(obj,fontname='微软雅黑',size=None):
    ## 设置字体函数
    obj.font.name = fontname
    obj._element.rPr.rFonts.set(qn('w:eastAsia'),fontname)
    if size and isinstance(size,Pt):
        obj.font.size = size

def gen_barcode(code='123456'):
    ## 生成二维码
    f = io.BytesIO()
    Code128Encoder(code).save(f)
    return f

def init_doc(doc):
    ## 设置页边距
    distance = Inches(0.3)
    sec = doc.sections[0]
    sec.left_margin = distance
    sec.right_margin = distance
    sec.top_margin = distance
    sec.bottom_margin = distance
    ##设置默认字体
    chg_font(doc.styles['Normal'],fontname='宋体')


def one_page(doc):

    layout_tab = doc.add_table(rows=4,cols=4)

    for cur_cell in (layout_tab.cell(0,0),layout_tab.cell(0,2),
            layout_tab.cell(1,0),layout_tab.cell(1,2),
            layout_tab.cell(2,0),layout_tab.cell(2,2),
            layout_tab.cell(3,0),layout_tab.cell(3,2),):
        ph = cur_cell.paragraphs[0]

        htitle = ph.add_run('2018年初中学业水平考试')
        htitle.add_break()
        title = ph.add_run('　准　考　证')
        title.add_break()
        chg_font(title,size=Pt(16))

        for left,right in zip(TEST_DATA,STUD_DATA):
            ph.add_run(left)
            run = ph.add_run(right)
            chg_font(run)
            run.bold = True
            run.add_break()

    for cur_cell in (layout_tab.cell(0,1),layout_tab.cell(0,3),
            layout_tab.cell(1,1),layout_tab.cell(1,3),
            layout_tab.cell(2,1),layout_tab.cell(2,3),
            layout_tab.cell(3,1),layout_tab.cell(3,3),):
        ph = cur_cell.paragraphs[0]
        run = ph.add_run()
        run.add_picture('aa.jpg',width=Inches(1.2))
        run.add_break()
        f = gen_barcode()
        run.add_picture(f,width=Inches(1.2))




if __name__ == '__main__':
    doc = Document()
    init_doc(doc)
    for i in range(50):
        one_page(doc)
        doc.add_page_break()
    doc.save('ab.docx')
