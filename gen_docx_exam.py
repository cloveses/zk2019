import os
import math
import io
from PIL import Image
from docx import Document
from docx.shared import Pt , Inches
from docx.oxml.ns import qn
from pystrich.code128 import Code128Encoder

from ph_models import *


## 生成准考证的WORD版DOCX，速度比生成PDF较快，体积稍小
## 每页8张准考证

TEST_DATA = ("准考号:  ","姓　名:  ","性　别:  ","考　点:  ","报名点:  ")
STUD_DATA = ('18251402',"李文娟",'男','泗县一中','01中学')

def confirm_path(path):
    ## 建立目录
    if not os.path.exists(path):
        os.makedirs(path)

def update_image(path):
    ## 修复问题照片文件（其它程序可查看和打开，此模块中报错）
    im = Image.open(path)
    im.save(path)
    im.close()

def chg_font(obj,fontname='微软雅黑',size=None):
    ## 设置字体函数
    obj.font.name = fontname
    obj._element.rPr.rFonts.set(qn('w:eastAsia'),fontname)
    if size and isinstance(size,Pt):
        obj.font.size = size

def gen_barcode(code='123456'):
    ## 生成二维码
    f = io.BytesIO()
    Code128Encoder(code).save(f)
    return f

def init_doc(doc):
    ## 设置页边距
    distance = Inches(0.3)
    sec = doc.sections[0]
    sec.left_margin = distance
    sec.right_margin = distance
    sec.top_margin = distance
    sec.bottom_margin = distance
    ##设置默认字体
    chg_font(doc.styles['Normal'],fontname='宋体')


def one_page(doc,studs):
    ## 生成一页内容（每页八张准考证）

    layout_tab = doc.add_table(rows=4,cols=4)

    for index,stud in enumerate(studs):
        r = index % 4
        c = (index // 4) * 2
        cur_cell = layout_tab.cell(r,c)
        ph = cur_cell.paragraphs[0]

        ## 插入考生信息
        htitle = ph.add_run('2018年初中学业水平考试')
        htitle.add_break()
        title = ph.add_run('　准　考　证')
        title.add_break()
        chg_font(title,size=Pt(16))

        for left,right in zip(TEST_DATA,stud[:5]):
            ph.add_run(left)
            run = ph.add_run(right)
            chg_font(run)
            run.bold = True
            run.add_break()

        ## 插入考生照片和条码图片
        cur_cell = layout_tab.cell(r,c+1)
        file_name = ''.join(('Z',stud[-1],'.jpg'))
        path = os.path.join('photos',stud[-2],file_name)
        ph = cur_cell.paragraphs[0]
        ph.add_run('　')
        run = ph.add_run()
        if os.path.exists(path):
            try:
                run.add_picture(path,width=Inches(1.2))
            except:
                # print(stud,'Photo file error!')
                # 如果照片错误，则重写照片文件
                update_image(path)
                run.add_picture(path,width=Inches(1.2))
        else:
            print(stud,'no photo file!')
        run.add_break()
        f = gen_barcode(stud[0])
        run.add_picture(f,width=Inches(1.5))

def gen_unit_docx(dir_name,sch_name,studs,page_num=8):
    ## 生成一个学校或一个分类的准考证
    confirm_path(dir_name)
    path = os.path.join(dir_name,sch_name + '.docx')
    pages = math.ceil(len(studs)/page_num)

    doc = Document()
    init_doc(doc)

    for i in range(pages):
        one_page(doc,studs[i*page_num:(i+1)*page_num])
        if i != pages - 1:
            doc.add_page_break()  #插入分页符

    doc.save(path)

# 按学校生成准考证
@db_session
def gen_examid_sch(dir_name):
    schs = select(s.sch for s in StudPh)
    # print(schs)
    # schs = ['泗县墩集中学',]
    for sch in schs:
        datas = select(s 
         for s in StudPh if s.sch==sch).order_by(StudPh.classcode,StudPh.phid)
        datas = [(s.phid,s.name,s.sex,s.exam_addr,s.sch,s.schcode,s.signid) for s in datas]
        gen_unit_docx(dir_name,sch,datas)


if __name__ == '__main__':
    gen_examid_sch('idsch')